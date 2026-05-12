"""
生成 5G 软件工程实训配套 Word 文档（简易商品库存管理系统）。
依赖：pip install python-docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def build():
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("软件工程实训 — 简易商品库存管理系统\n文档与报告（纲要）")
    run.bold = True
    run.font.size = Pt(16)

    add_para(doc, "说明：下文为可直接提交的纲要；请按教师要求补充封面、个人信息、截图、签字及日期。")

    add_heading(doc, "一、需求分析报告", 1)
    add_heading(doc, "1. 项目背景与目标", 2)
    add_para(
        doc,
        "背景：小型商户或实训场景需要对商品编号、名称、类别、库存数量与单价进行统一管理，"
        "替代纸质或零散表格，降低出错率并提高查询效率。\n"
        "目标：实现基于 Web 的简易商品库存管理系统，支持用户注册登录与商品的增删改查，"
        "数据持久化保存，软件重启后数据不丢失。",
    )
    add_heading(doc, "2. 功能需求", 2)
    add_para(
        doc,
        "（1）登录注册模块：用户名/密码非空校验；注册查重；登录成功后进入主界面。\n"
        "（2）核心业务（库存）：新增商品、删除（确认）、修改、列表展示；"
        "查询支持关键词，对编号、名称、类别、备注进行模糊匹配。\n"
        "（3）系统导航：主界面菜单包含库存列表、新增商品、退出登录。",
    )
    add_heading(doc, "3. 非功能需求", 2)
    add_para(
        doc,
        "界面：简洁清晰的中文界面，表格展示库存；错误与成功提示明确。\n"
        "性能：单机小型数据量（数百～数千条）下页面响应及时。\n"
        "数据存储：采用 SQLite 嵌入式数据库，文件持久化；必要时可备份 data/inventory.db。\n"
        "安全：密码使用哈希存储；会话基于服务端 Session（实训演示环境）。",
    )
    add_heading(doc, "4. 运行环境说明", 2)
    add_para(
        doc,
        "开发与运行：Windows 10/11，Python 3.10+，Flask 3.x；浏览器访问 http://127.0.0.1:5000 。\n"
        "依赖安装：pip install -r requirements.txt；初始化数据库：python -m flask --app app init-db 。",
    )

    add_heading(doc, "二、总体设计", 1)
    add_heading(doc, "1. 模块划分", 2)
    add_para(
        doc,
        "（1）认证模块：注册、登录校验、退出、会话管理。\n"
        "（2）库存业务模块：商品 CRUD、条件查询、表单校验。\n"
        "（3）数据模块：SQLite 访问层、建表脚本 schema.sql、初始化命令。",
    )
    add_heading(doc, "2. 数据库设计（逻辑结构）", 2)
    add_para(
        doc,
        "实体 users：用户（id，username 唯一，password_hash，created_at）。\n"
        "实体 products：商品（id，sku 唯一，name，category，quantity，unit_price，note，created_at，updated_at）。\n"
        "关系：用户与商品在本实训中为独立实体；业务上由已登录用户操作系统维护商品数据（无外键约束到 users，简化实现）。",
    )
    add_para(
        doc,
        "E-R 图说明（请在报告中手绘或使用工具绘制）：\n"
        "• USER ──< 操作 >── PRODUCT（多对多简化为：登录用户对 PRODUCT 表执行维护操作）。\n"
        "• PRODUCT 关键属性：SKU 唯一标识库存条目。",
    )
    add_heading(doc, "3. 系统架构", 2)
    add_para(
        doc,
        "采用 B/S 结构：浏览器 ↔ Flask（Python） ↔ SQLite 文件。\n"
        "请求流程：路由接收表单 → 校验 → SQL 执行 → 渲染模板返回 HTML。",
    )

    add_heading(doc, "三、详细设计", 1)
    add_heading(doc, "1. 登录校验流程", 2)
    add_para(
        doc,
        "接收用户名密码 → 非空判断 → 查询 users 表 → Werkzeug 校验密码哈希 → "
        "成功则写入 session(user_id, username) 并重定向至库存列表；失败提示错误。",
    )
    add_heading(doc, "2. 商品新增 / 修改 / 删除 / 查询", 2)
    add_para(
        doc,
        "新增：表单字段校验（SKU 格式、数量非负整数、单价非负）→ INSERT；SKU 冲突返回提示。\n"
        "修改：按 id 加载原记录 → 校验 → UPDATE 并刷新 updated_at。\n"
        "删除：POST 提交 → DELETE WHERE id → 列表刷新。\n"
        "查询：可选参数 q → SQL LIKE 模糊匹配多字段 → ORDER BY updated_at DESC。",
    )
    add_heading(doc, "3. 界面与流程", 2)
    add_para(
        doc,
        "登录页、注册页、库存列表（检索框 + 表格 + 操作列）、商品表单页（新增/编辑）。"
        "删除前浏览器 confirm 二次确认。",
    )

    add_heading(doc, "四、编码实现说明", 1)
    add_para(
        doc,
        "技术选型：Flask + Jinja2 模板 + SQLite。\n"
        "源码目录：inventory_system/（app.py、db.py、templates、static、schema.sql）。\n"
        "启动：python app.py；首次注册账号后即可录入商品数据。",
    )

    add_heading(doc, "五、系统测试报告", 1)
    add_heading(doc, "1. 测试用例（示例）", 2)
    add_para(
        doc,
        "| 编号 | 模块 | 步骤 | 预期结果 |\n"
        "| TC-L01 | 登录 | 用户名为空提交 | 提示不能为空，停留登录页 |\n"
        "| TC-L02 | 登录 | 错误密码 | 提示用户名或密码错误 |\n"
        "| TC-R01 | 注册 | 两次密码不一致 | 提示不一致 |\n"
        "| TC-P01 | 商品 | 新增合法 SKU | 列表出现新记录 |\n"
        "| TC-P02 | 商品 | SKU 重复 | 提示编号已存在 |\n"
        "| TC-P03 | 商品 | 关键词查询 | 命中记录展示 |\n"
        "| TC-P04 | 商品 | 修改库存与单价 | 列表数值更新 |\n"
        "| TC-P05 | 商品 | 删除并确认 | 记录移除 |\n"
        "| TC-D01 | 持久化 | 重启服务后登录 | 历史数据仍在 |\n"
        "（请在正式报告中将「实际结果」列补充为通过/失败，并附截图。）",
    )
    add_heading(doc, "2. 自动化测试", 2)
    add_para(
        doc,
        "项目包含 pytest 用例：tests/test_app.py；命令 pytest -q。"
        "覆盖注册登录后的新增、查询、修改、删除基本路径。",
    )

    add_heading(doc, "六、实训报告（总结提纲）", 1)
    add_para(
        doc,
        "1）本人在需求分析与设计阶段的理解与分工说明；\n"
        "2）编码阶段遇到的数据校验、会话管理等问题与解决思路；\n"
        "3）测试阶段发现的缺陷与回归结果；\n"
        "4）收获与改进方向（如权限细分、入库出库单据、条码扫描等可扩展点）。",
    )

    out = Path(__file__).resolve().parent / "软件工程实训_简易商品库存管理系统_文档.docx"
    doc.save(out)
    print("已生成:", out)


if __name__ == "__main__":
    build()
